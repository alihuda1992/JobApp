"""
Helpers for editing Ali's master resume template in place, preserving its
exact formatting (fonts, bold/italic/underline, tab stops, bullet style)
rather than rebuilding a resume from scratch.

Import from render_resume.py or use directly:

    from build_resume import load, set_paragraph_text, find_paragraph,
        replace_bullets_after, remove_role_block, set_skill_category,
        purge_empty_bullets, save
"""
import copy
import docx
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn


def load(template_path):
    return docx.Document(template_path)


def _copy_run_format(src_run, dst_run):
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    if src_run.font.size:
        dst_run.font.size = src_run.font.size
    if src_run.font.name:
        dst_run.font.name = src_run.font.name


def set_paragraph_text(paragraph, text, style_run=None):
    """Replace a paragraph's visible text with `text`, keeping paragraph-level
    formatting. Copies character formatting from `style_run` (defaults to the
    paragraph's own first run)."""
    if style_run is None:
        if not paragraph.runs:
            raise ValueError("Paragraph has no runs to copy formatting from — pass style_run explicitly")
        style_run = paragraph.runs[0]
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    new_run = paragraph.add_run(text)
    _copy_run_format(style_run, new_run)
    return new_run


def find_paragraph(doc, text_startswith, occurrence=1):
    """Find the Nth paragraph (1-indexed) whose text starts with `text_startswith`.
    Raises ValueError if not found."""
    count = 0
    for p in doc.paragraphs:
        if p.text.strip().startswith(text_startswith):
            count += 1
            if count == occurrence:
                return p
    raise ValueError(f"Paragraph starting with {text_startswith!r} (occurrence {occurrence}) not found")


def _is_bullet(paragraph):
    if paragraph.style.name == 'List Paragraph':
        return True
    pPr = paragraph._element.pPr
    if pPr is None:
        return False
    return pPr.find(qn('w:numPr')) is not None


def replace_bullets_after(doc, anchor_paragraph, new_bullets):
    """Replace the bullet paragraphs immediately following anchor_paragraph
    with new_bullets. Pass an empty list to remove all bullets."""
    anchor_el = anchor_paragraph._element
    existing = []
    el = anchor_el.getnext()
    while el is not None and el.tag == qn('w:p'):
        p = Paragraph(el, anchor_paragraph._parent)
        if not _is_bullet(p):
            break
        existing.append(p)
        el = el.getnext()

    if not existing and new_bullets:
        raise ValueError(
            "No bullet paragraphs found after {!r} — "
            "check anchor text against resume-template-spec.md".format(anchor_paragraph.text[:60])
        )

    template_bullet = existing[0] if existing else None
    style_run = template_bullet.runs[0] if template_bullet and template_bullet.runs else None

    for i, bullet_text in enumerate(new_bullets):
        if i < len(existing):
            set_paragraph_text(existing[i], bullet_text, style_run=style_run)
        else:
            new_p_el = copy.deepcopy(template_bullet._element)
            existing[-1]._element.addnext(new_p_el)
            new_p = Paragraph(new_p_el, anchor_paragraph._parent)
            set_paragraph_text(new_p, bullet_text, style_run=style_run)
            existing.append(new_p)

    for extra in existing[len(new_bullets):]:
        extra._element.getparent().remove(extra._element)


def remove_paragraph(paragraph):
    """Delete a paragraph entirely."""
    paragraph._element.getparent().remove(paragraph._element)


def remove_role_block(doc, anchor_paragraph):
    """Remove an entire role block — clears bullets first, then the anchor
    (role/date) line. Call remove_paragraph() separately on the company-name
    line if the whole company block is being dropped."""
    replace_bullets_after(doc, anchor_paragraph, [])
    remove_paragraph(anchor_paragraph)


def set_skill_category(doc, label, items, occurrence=1):
    """Update a Skills section line that has a bold label run and a plain
    items run (e.g. 'Transformation:  operating model design, AI adoption').
    `label` should include the trailing colon and two spaces."""
    p = find_paragraph(doc, label.rstrip(": "), occurrence)
    if len(p.runs) < 2:
        raise ValueError(f"Skills paragraph {label!r} expected 2 runs (bold label + plain items), got {len(p.runs)}")
    # Run 0: bold label — preserve bold, just rewrite text
    p.runs[0].text = label
    # Run 1: plain items
    p.runs[1].text = items


def purge_empty_bullets(doc):
    """Remove bullet-style paragraphs with no text — template leftovers that
    push content to a second page. Call once after all edits, before save()."""
    to_remove = [p for p in doc.paragraphs if _is_bullet(p) and not p.text.strip()]
    for p in to_remove:
        p._element.getparent().remove(p._element)
    return len(to_remove)


def save(doc, out_path):
    doc.save(out_path)
